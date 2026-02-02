"""
==============================================================================
WIZARD MODE - Interactive Plan Review Checklist
==============================================================================
File: pages/2_Wizard_Mode.py
Purpose: Guide reviewers through plan review with interactive checklists
         and automatic comment generation.

Features:
- Review type selection (Transitional, HP, Standard, Pool, Fence)
- Interactive Yes/No/N/A radio buttons per checklist item
- Automatic comment suggestions when "No" is selected
- Custom notes support
- Word document export with full checklist and comments
- LAMA CSV export for LAMA Comment Uploader chrome extension
- Bluebeam BAX export with full styling (green text boxes, Helvetica 12pt)

Update Log:
- 2026-02-02: Added explicit text colors and !important to all CSS classes
              to fix invisible text when browser is in dark mode.
- 2026-02-02: Added custom sidebar navigation to replace default "app" label.
- 2026-02-02: Added LAMA CSV and Bluebeam BAX export buttons.
- 2026-02-02: Changed status dropdowns to horizontal radio buttons for speed.
- 2026-02-02: Switched from XML MarkupSummary to BAX (Bluebeam Markup Archive)
              format which carries full annotation styling in the Raw field
              including green border/fill, 25% fill opacity, Helvetica 12pt,
              and 0.75pt solid border.
==============================================================================
"""

import streamlit as st
import sys
import csv
import zlib
import random
import string
from pathlib import Path
from datetime import datetime
from io import BytesIO, StringIO

# Add utils to path
sys.path.append(str(Path(__file__).parent.parent / "utils"))

from checklist_data import (
    REVIEW_TYPES, 
    REVIEWERS, 
    CHECKLIST_SECTIONS,
    get_checklist_for_review_type
)
from comments_database import COMMENTS, get_comment

# Import python-docx for Word export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

st.set_page_config(page_title="Wizard Mode", page_icon="📋", layout="wide")

# =============================================================================
# CUSTOM CSS - Dark Mode Safe
# =============================================================================
# Every class with a light background explicitly sets color: #1a1a2e (dark navy)
# with !important so Streamlit's dark theme cannot override text to white.
# Radio button styling keeps items compact and inline.
# =============================================================================
st.markdown("""
<style>
    /* Hide default streamlit page navigation */
    [data-testid="stSidebarNav"] {
        display: none;
    }
    .review-header {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%) !important;
        color: white !important;
        padding: 1.5rem;
        border-radius: 10px;
        margin-bottom: 1rem;
    }
    .section-header {
        background: #f0f4f8 !important;
        color: #1a1a2e !important;
        padding: 0.75rem 1rem;
        border-left: 4px solid #1e3a5f;
        margin: 1rem 0 0.5rem 0;
        font-weight: bold;
    }
    .checklist-item {
        background: #ffffff !important;
        color: #1a1a2e !important;
        border: 1px solid #e0e0e0;
        border-radius: 5px;
        padding: 0.75rem;
        margin: 0.5rem 0;
    }
    .comment-box {
        background: #fff8e1 !important;
        color: #1a1a2e !important;
        border: 1px solid #ffd54f;
        border-radius: 5px;
        padding: 0.75rem;
        margin: 0.5rem 0 0.5rem 1rem;
        font-size: 0.9em;
    }
    .status-yes { color: #2e7d32 !important; font-weight: bold; }
    .status-no { color: #c62828 !important; font-weight: bold; }
    .status-na { color: #757575 !important; font-weight: bold; }
    .export-section {
        background: #e8f5e9 !important;
        color: #1a1a2e !important;
        border: 1px solid #a5d6a7;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
    /* Compact radio buttons - reduce vertical padding */
    div[data-testid="stRadio"] > div {
        gap: 0.5rem;
    }
</style>
""", unsafe_allow_html=True)

# Custom Sidebar Navigation (matches app.py)
st.sidebar.title("🧭 Navigation")
st.sidebar.markdown("---")
st.sidebar.page_link("app.py", label="🏡 Dashboard", icon=None)
st.sidebar.page_link("pages/1_QA_Mode.py", label="💬 Q&A Mode", icon=None)
st.sidebar.page_link("pages/2_Wizard_Mode.py", label="🧙‍♂️ Wizard Mode", icon=None)
st.sidebar.markdown("---")
st.sidebar.markdown("**Engineering AI Assistant**")
st.sidebar.markdown("v1.0 | Brentwood, TN")


def initialize_session_state():
    """Initialize session state variables"""
    if 'wizard_review_type' not in st.session_state:
        st.session_state.wizard_review_type = None
    if 'wizard_permit_number' not in st.session_state:
        st.session_state.wizard_permit_number = ""
    if 'wizard_address' not in st.session_state:
        st.session_state.wizard_address = ""
    if 'wizard_reviewer' not in st.session_state:
        st.session_state.wizard_reviewer = None
    if 'wizard_checklist_state' not in st.session_state:
        st.session_state.wizard_checklist_state = {}
    if 'wizard_selected_comments' not in st.session_state:
        st.session_state.wizard_selected_comments = {}
    if 'wizard_custom_notes' not in st.session_state:
        st.session_state.wizard_custom_notes = {}
    if 'wizard_started' not in st.session_state:
        st.session_state.wizard_started = False


def reset_checklist():
    """Reset checklist state when review type changes"""
    st.session_state.wizard_checklist_state = {}
    st.session_state.wizard_selected_comments = {}
    st.session_state.wizard_custom_notes = {}


def collect_all_comments():
    """
    Collect all comments from checklist items marked 'No'.
    Returns a list of raw comment strings (no prefix codes).
    Used by all export functions to ensure consistent output.
    """
    checklist = get_checklist_for_review_type(st.session_state.wizard_review_type)
    all_comments = []

    for section_id, section_data in checklist.items():
        for item in section_data["items"]:
            item_key = item["id"]
            if st.session_state.wizard_checklist_state.get(item_key) == "No":
                selected = st.session_state.wizard_selected_comments.get(item_key, [])
                custom_note = st.session_state.wizard_custom_notes.get(item_key, "")

                for comment_id in selected:
                    comment_text = COMMENTS.get(comment_id, "")
                    if comment_text:
                        all_comments.append(comment_text)

                if custom_note.strip():
                    all_comments.append(custom_note.strip())

    return all_comments


def generate_word_document():
    """Generate a Word document with review comments"""
    if not DOCX_AVAILABLE:
        return None
    
    doc = Document()
    
    # =========================================================================
    # TITLE
    # =========================================================================
    title = doc.add_heading('Engineering Plan Review', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =========================================================================
    # PROJECT INFORMATION
    # =========================================================================
    doc.add_heading('Project Information', level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Review Type:', st.session_state.wizard_review_type or 'Not specified'),
        ('Permit Number:', st.session_state.wizard_permit_number or 'Not specified'),
        ('Address:', st.session_state.wizard_address or 'Not specified'),
        ('Reviewer:', st.session_state.wizard_reviewer or 'Not specified'),
        ('Review Date:', datetime.now().strftime('%Y-%m-%d %H:%M')),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        # Bold the labels
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # =========================================================================
    # SUMMARY STATISTICS
    # =========================================================================
    doc.add_heading('Review Summary', level=1)
    
    yes_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "Yes")
    no_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "No")
    na_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "N/A")
    total = yes_count + no_count + na_count
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Total Items Reviewed:', str(total)),
        ('Compliant (Yes):', str(yes_count)),
        ('Issues Found (No):', str(no_count)),
        ('Not Applicable:', str(na_count)),
    ]
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # =========================================================================
    # FULL CHECKLIST WITH STATUS AND COMMENTS
    # =========================================================================
    doc.add_heading('Plan Review Checklist', level=1)
    
    # Get the checklist for the current review type
    checklist = get_checklist_for_review_type(st.session_state.wizard_review_type)
    
    for section_id, section_data in checklist.items():
        # Section header
        section_heading = doc.add_heading(section_data["name"], level=2)
        
        for item in section_data["items"]:
            item_key = item["id"]
            status = st.session_state.wizard_checklist_state.get(item_key, "Not Reviewed")
            
            # Create paragraph for checklist item
            para = doc.add_paragraph()
            
            # Add item ID and description
            item_run = para.add_run(f"{item['id']} - {item['description']}")
            item_run.bold = False
            
            # Add status on new line with color/bold
            para.add_run("\n")
            status_run = para.add_run(f"Status: {status}")
            status_run.bold = True
            
            # Color code the status
            if status == "Yes":
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif status == "No":
                status_run.font.color.rgb = RGBColor(200, 0, 0)  # Red
            elif status == "N/A":
                status_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
            else:
                status_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange for not reviewed
            
            # If "No", show comments
            if status == "No":
                selected = st.session_state.wizard_selected_comments.get(item_key, [])
                custom_note = st.session_state.wizard_custom_notes.get(item_key, "")
                
                if selected or custom_note.strip():
                    para.add_run("\n")
                    comments_label = para.add_run("Comments:")
                    comments_label.bold = True
                    comments_label.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
                    
                    # Add each selected comment
                    for comment_id in selected:
                        comment_text = COMMENTS.get(comment_id, "")
                        if comment_text:
                            comment_para = doc.add_paragraph()
                            comment_para.paragraph_format.left_indent = Inches(0.5)
                            comment_run = comment_para.add_run(f"• [{comment_id}] {comment_text}")
                            comment_run.font.size = Pt(10)
                    
                    # Add custom note if provided
                    if custom_note.strip():
                        custom_para = doc.add_paragraph()
                        custom_para.paragraph_format.left_indent = Inches(0.5)
                        custom_run = custom_para.add_run(f"• [CUSTOM] {custom_note}")
                        custom_run.font.size = Pt(10)
                        custom_run.italic = True
        
        # Add space after each section
        doc.add_paragraph()
    
    # =========================================================================
    # COMMENTS FOR COPY/PASTE (Only "No" items)
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('Comments for Copy/Paste', level=1)
    
    intro_para = doc.add_paragraph()
    intro_para.add_run("Use the comments below for Bluebeam or permit system. ").italic = True
    intro_para.add_run("Only items marked 'No' with selected comments are included.").italic = True
    doc.add_paragraph()
    
    # Collect all comments in order
    all_comments = []
    
    for section_id, section_data in checklist.items():
        for item in section_data["items"]:
            item_key = item["id"]
            if st.session_state.wizard_checklist_state.get(item_key) == "No":
                selected = st.session_state.wizard_selected_comments.get(item_key, [])
                custom_note = st.session_state.wizard_custom_notes.get(item_key, "")
                
                for comment_id in selected:
                    comment_text = COMMENTS.get(comment_id, "")
                    if comment_text:
                        all_comments.append(comment_text)
                
                if custom_note.strip():
                    all_comments.append(custom_note)
    
    if all_comments:
        for i, comment in enumerate(all_comments, 1):
            para = doc.add_paragraph()
            num_run = para.add_run(f"{i}. ")
            num_run.bold = True
            para.add_run(comment)
            doc.add_paragraph()  # Space between comments
    else:
        doc.add_paragraph("No comments to include - all items are compliant or N/A.")
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# LAMA CSV EXPORT
# =============================================================================
# Generates a single-column CSV with header "Comments" that plugs directly
# into the LAMA Comment Uploader chrome extension. The extension's panel.js
# parser looks for a column header matching "comments" (case-insensitive)
# and reads each row as a comment to post to LAMA.
# =============================================================================

def generate_lama_csv():
    """
    Generate CSV for the LAMA Comment Uploader chrome extension.
    Format: Single column with header 'Comments', RFC 4180 quoting.
    """
    comments = collect_all_comments()
    if not comments:
        return None

    buffer = StringIO()
    writer = csv.writer(buffer, quoting=csv.QUOTE_ALL)
    writer.writerow(["Comments"])
    for comment in comments:
        writer.writerow([comment])

    return buffer.getvalue().encode('utf-8')


# =============================================================================
# BLUEBEAM BAX EXPORT (Bluebeam Markup Archive)
# =============================================================================
# Generates a BAX file — Bluebeam Revu's native markup archive format.
# Import via: Markup → Import → select .bax file
#
# Unlike the XML MarkupSummary format, BAX files contain a <Raw> field
# with a zlib-compressed PDF annotation dictionary that carries FULL
# visual styling. This means imported markups arrive with the exact
# appearance from Kevin's toolchest:
#
#   Border color:   #008000 (green)  →  C[0 0.5019608 0]
#   Fill color:     #008000 (green)  →  DA: 0 0.5019608 0 rg
#   Fill opacity:   25%              →  FillOpacity 0.25
#   Border width:   0.75 pt, solid   →  BS<</W 0.75/S/S>>
#   Font:           Helvetica 12pt   →  /Helv 12 Tf
#   Text color:     Black            →  color:#000000
#   Line height:    13.8pt           →  line-height:13.8pt
#   Alignment:      Left, Top        →  text-align:left
#
# The Raw field structure was reverse-engineered from Kevin's actual
# Bluebeam Revu 21.8 export (Test.bax, 2026-02-02). Each Raw field
# decompresses to a PDF annotation dictionary like:
#   <</DA(...)/DS(...)/Rect[...]/Contents(...)/RC(...)/...>>
#
# Annotations are placed on page 1, stacked vertically with overflow
# to a second column. User repositions them onto the correct sheets
# after import.
# =============================================================================

def _generate_annotation_id():
    """
    Generate a 16-character uppercase letter ID matching Bluebeam's
    annotation naming convention (e.g., 'QYFOJAMUAUFKGRPK').
    """
    return ''.join(random.choices(string.ascii_uppercase, k=16))


def _pdf_escape(text):
    """
    Escape special characters for PDF string literals inside ().
    In PDF, only backslash, open-paren, and close-paren need escaping.
    """
    return (str(text)
            .replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)"))


def _xml_escape(text):
    """Escape special characters for XML text content."""
    return (str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&apos;"))


def _build_annotation_raw(comment_text, reviewer, annot_id, rect, pdf_date):
    """
    Build the zlib-compressed, hex-encoded Raw field for a BAX annotation.

    This constructs a PDF annotation dictionary string matching the exact
    format produced by Bluebeam Revu 21.8, then compresses it with zlib
    and returns the hex-encoded result.

    Parameters:
        comment_text: The plain text comment
        reviewer:     Author name (from wizard reviewer dropdown)
        annot_id:     16-char unique annotation ID
        rect:         Tuple of (x1, y1, x2, y2) in PDF points
        pdf_date:     Date string in PDF format: D:YYYYMMDDHHMMSS-06'00'

    Returns:
        Hex string of zlib-compressed PDF annotation dictionary
    """
    x1, y1, x2, y2 = rect

    # Escape text for PDF string literals (inside parentheses)
    pdf_text = _pdf_escape(comment_text)
    pdf_reviewer = _pdf_escape(reviewer)

    # For RC (rich content): HTML-escape first, then PDF-escape the result
    # because the HTML sits inside a PDF string literal
    html_text = (comment_text
                 .replace("&", "&amp;")
                 .replace("<", "&lt;")
                 .replace(">", "&gt;"))
    rc_text = _pdf_escape(html_text)

    # -------------------------------------------------------------------------
    # Build the PDF annotation dictionary
    # This matches the exact structure from Kevin's Bluebeam Revu 21.8 export
    # -------------------------------------------------------------------------
    raw_str = (
        # Default Appearance: green fill (0 0.5019608 0), Helvetica 12pt
        '<</DA(0 0.5019608 0 rg /Helv 12 Tf)'
        # Default Style: CSS-like styling for Bluebeam's text renderer
        '/DS(font: Helvetica 12pt; text-align:left; margin:0pt; '
        'line-height:13.8pt; color:#000000)'
        # Temporary bounding box (matches Rect)
        f'/TempBBox[{x1} {y1} {x2} {y2}]'
        # Fill opacity: 25% (green background at 25% transparency)
        '/FillOpacity 0.25'
        # Title = Author name
        f'/T({pdf_reviewer})'
        # Creation date in PDF format
        f'/CreationDate({pdf_date})'
        # Rich Content: XHTML body with styled text
        '/RC(<?xml version="1.0"?>'
        '<body xmlns:xfa="http://www.xfa.org/schema/xfa-data/1.0/"'
        ' xfa:contentType="text/html"'
        ' xfa:APIVersion="BluebeamPDFRevu:2018"'
        ' xfa:spec="2.2.0"'
        ' style="font: Helvetica 12pt; text-align:left; margin:0pt; '
        'line-height:13.8pt; color:#000000"'
        ' xmlns="http://www.w3.org/1999/xhtml">'
        f'<p>{rc_text}</p></body>)'
        # Subject category
        '/Subj(Engineering)'
        # Unique annotation name
        f'/NM({annot_id})'
        # Annotation subtype
        '/Subtype/FreeText'
        # Rectangle coordinates [x1 y1 x2 y2] in PDF points
        f'/Rect[{x1} {y1} {x2} {y2}]'
        # Plain text content (fallback for non-rich-text readers)
        f'/Contents({pdf_text})'
        # Flags: 4 = Print (annotation prints with document)
        '/F 4'
        # Color array: green (0, 0.5019608, 0) = #008000
        '/C[0 0.5019608 0]'
        # Border style: 0.75pt solid
        '/BS<</W 0.75/S/S/Type/Border>>'
        # Modified date
        f'/M({pdf_date})>>'
    )

    # Compress with zlib (produces 789c... header matching Bluebeam's output)
    compressed = zlib.compress(raw_str.encode('utf-8'))
    return compressed.hex()


def generate_bluebeam_bax():
    """
    Generate a Bluebeam BAX (Markup Archive) file with fully styled
    FreeText annotations matching Kevin's toolchest configuration.

    Returns bytes of the complete BAX file (UTF-8 with BOM, CRLF endings)
    or None if there are no comments to export.
    """
    comments = collect_all_comments()
    if not comments:
        return None

    reviewer = st.session_state.wizard_reviewer or "Engineering"
    now = datetime.now()

    # -------------------------------------------------------------------------
    # Date formats
    # Outer XML uses ISO 8601 with .0000000Z suffix (Bluebeam convention)
    # Inner PDF uses D:YYYYMMDDHHMMSS-06'00' (Central Time offset)
    # -------------------------------------------------------------------------
    iso_date = now.strftime("%Y-%m-%dT%H:%M:%S") + ".0000000Z"
    pdf_date = now.strftime("D:%Y%m%d%H%M%S") + "-06'00'"

    # -------------------------------------------------------------------------
    # Page and annotation layout
    # All values in PDF points (1 inch = 72 points)
    # US Letter: 612 x 792 points (8.5" x 11")
    # -------------------------------------------------------------------------
    page_width = 612
    page_height = 792
    box_width = 252     # 3.5 inches
    box_height = 108    # 1.5 inches
    margin = 36         # 0.5 inch margin from edges
    gap = 10            # ~0.14 inch vertical gap between stacked boxes

    # Start stacking from top of page, moving downward
    # PDF origin is bottom-left, Y increases upward
    current_y_top = page_height - margin
    x1 = margin

    # -------------------------------------------------------------------------
    # Build annotation XML elements
    # -------------------------------------------------------------------------
    annotation_blocks = []

    for i, comment in enumerate(comments):
        y2 = current_y_top                 # Top of box
        y1 = current_y_top - box_height    # Bottom of box

        # If we've gone below the bottom margin, start a second column
        if y1 < margin and x1 == margin:
            x1 = margin + box_width + margin
            current_y_top = page_height - margin
            y2 = current_y_top
            y1 = current_y_top - box_height

        x2 = x1 + box_width
        rect = (x1, y1, x2, y2)

        # Generate unique annotation ID
        annot_id = _generate_annotation_id()

        # Build the compressed Raw field with full styling
        raw_hex = _build_annotation_raw(comment, reviewer, annot_id, rect, pdf_date)

        # Build outer XML for this annotation
        annotation_xml = (
            '    <Annotation>\n'
            '      <Page>1</Page>\n'
            f'      <Contents>{_xml_escape(comment)}</Contents>\n'
            f'      <ModDate>{iso_date}</ModDate>\n'
            '      <Color>#008000</Color>\n'
            '      <Type>FreeText</Type>\n'
            f'      <ID>{annot_id}</ID>\n'
            '      <TypeInternal>Bluebeam.PDF.Annotations.AnnotationFreeText'
            '</TypeInternal>\n'
            f'      <Raw>{raw_hex}</Raw>\n'
            f'      <Index>{i}</Index>\n'
            '      <Subject>Engineering</Subject>\n'
            f'      <CreationDate>{iso_date}</CreationDate>\n'
            f'      <Author>{_xml_escape(reviewer)}</Author>\n'
            '    </Annotation>'
        )

        annotation_blocks.append(annotation_xml)
        current_y_top = y1 - gap

    # -------------------------------------------------------------------------
    # Assemble complete BAX document
    # -------------------------------------------------------------------------
    annotations_str = '\n'.join(annotation_blocks)

    bax_content = (
        '<?xml version="1.0" encoding="utf-8"?>\n'
        '<Document Version="1">\n'
        '  <Page Index="0">\n'
        '    <Label>1</Label>\n'
        f'    <Width>{page_width}</Width>\n'
        f'    <Height>{page_height}</Height>\n'
        f'{annotations_str}\n'
        '  </Page>\n'
        '</Document>'
    )

    # Encode with UTF-8 BOM and CRLF line endings to match Bluebeam's format
    bax_crlf = bax_content.replace('\n', '\r\n')
    return b'\xef\xbb\xbf' + bax_crlf.encode('utf-8')


def main():
    """Main function for Wizard Mode"""
    initialize_session_state()
    
    st.title("📋 Engineering Review Wizard")
    st.markdown("Interactive checklist for plan reviews with automatic comment generation.")
    
    # =========================================================================
    # STEP 1: PROJECT SETUP
    # =========================================================================
    st.markdown("---")
    st.subheader("📝 Step 1: Project Setup")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        review_type = st.selectbox(
            "Review Type",
            options=[""] + REVIEW_TYPES,
            index=0 if not st.session_state.wizard_review_type else REVIEW_TYPES.index(st.session_state.wizard_review_type) + 1,
            key="review_type_select"
        )
        
        # Check if review type changed
        if review_type and review_type != st.session_state.wizard_review_type:
            st.session_state.wizard_review_type = review_type
            reset_checklist()
            st.rerun()
        elif review_type:
            st.session_state.wizard_review_type = review_type
    
    with col2:
        permit_number = st.text_input(
            "Permit Number",
            value=st.session_state.wizard_permit_number,
            placeholder="e.g., SW2024-001"
        )
        st.session_state.wizard_permit_number = permit_number
    
    with col3:
        address = st.text_input(
            "Address",
            value=st.session_state.wizard_address,
            placeholder="e.g., 1808 Sonoma Trce"
        )
        st.session_state.wizard_address = address
    
    with col4:
        reviewer = st.selectbox(
            "Reviewer",
            options=[""] + REVIEWERS,
            index=0 if not st.session_state.wizard_reviewer else REVIEWERS.index(st.session_state.wizard_reviewer) + 1
        )
        st.session_state.wizard_reviewer = reviewer if reviewer else None
    
    # Check if we can proceed
    if not st.session_state.wizard_review_type:
        st.info("👆 Select a review type to begin.")
        return
    
    # =========================================================================
    # STEP 2: INTERACTIVE CHECKLIST
    # =========================================================================
    st.markdown("---")
    st.subheader(f"📋 Step 2: {st.session_state.wizard_review_type} Checklist")
    
    # Get applicable checklist
    checklist = get_checklist_for_review_type(st.session_state.wizard_review_type)
    
    # Calculate progress
    total_items = sum(len(section["items"]) for section in checklist.values())
    completed_items = len(st.session_state.wizard_checklist_state)
    
    st.progress(completed_items / total_items if total_items > 0 else 0)
    st.caption(f"Progress: {completed_items} of {total_items} items reviewed")
    
    # Display checklist by section
    for section_id, section_data in checklist.items():
        st.markdown(f'<div class="section-header">{section_data["name"]}</div>', unsafe_allow_html=True)
        
        for item in section_data["items"]:
            item_key = item["id"]
            
            # Create columns: description on left, radio buttons on right
            col1, col2 = st.columns([3, 1])
            
            with col1:
                st.markdown(f"**{item['id']}** - {item['description']}")
            
            with col2:
                # -------------------------------------------------------
                # Horizontal radio buttons instead of dropdown
                # Options: "—" (not yet reviewed), "Yes", "No", "N/A"
                # The "—" option acts as a blank/unreviewed placeholder
                # -------------------------------------------------------
                current_status = st.session_state.wizard_checklist_state.get(item_key, "—")
                
                status = st.radio(
                    f"Status for {item_key}",
                    options=["—", "Yes", "No", "N/A"],
                    index=["—", "Yes", "No", "N/A"].index(current_status) if current_status in ["—", "Yes", "No", "N/A"] else 0,
                    key=f"status_{item_key}",
                    horizontal=True,
                    label_visibility="collapsed"
                )
                
                # Update session state based on selection
                if status and status != "—":
                    st.session_state.wizard_checklist_state[item_key] = status
                elif item_key in st.session_state.wizard_checklist_state:
                    del st.session_state.wizard_checklist_state[item_key]
            
            # If "No" is selected, show comment options
            if st.session_state.wizard_checklist_state.get(item_key) == "No":
                with st.container():
                    st.markdown('<div class="comment-box">', unsafe_allow_html=True)
                    st.markdown("**📝 Select applicable comments:**")
                    
                    # Get available comments for this item
                    comment_ids = item.get("comment_ids", [])
                    
                    if comment_ids:
                        # Initialize selected comments if not exists
                        if item_key not in st.session_state.wizard_selected_comments:
                            st.session_state.wizard_selected_comments[item_key] = []
                        
                        # Show each comment with checkbox
                        for comment_id in comment_ids:
                            comment_text = COMMENTS.get(comment_id, "Comment not found")
                            
                            # Truncate long comments for display
                            display_text = comment_text[:150] + "..." if len(comment_text) > 150 else comment_text
                            
                            is_selected = comment_id in st.session_state.wizard_selected_comments[item_key]
                            
                            if st.checkbox(
                                f"**{comment_id}**: {display_text}",
                                value=is_selected,
                                key=f"comment_{item_key}_{comment_id}"
                            ):
                                if comment_id not in st.session_state.wizard_selected_comments[item_key]:
                                    st.session_state.wizard_selected_comments[item_key].append(comment_id)
                            else:
                                if comment_id in st.session_state.wizard_selected_comments[item_key]:
                                    st.session_state.wizard_selected_comments[item_key].remove(comment_id)
                            
                            # Show full comment in expander if long
                            if len(comment_text) > 150:
                                with st.expander("View full comment"):
                                    st.write(comment_text)
                    
                    # Custom notes field
                    st.markdown("**✏️ Custom notes (optional):**")
                    custom_note = st.text_area(
                        "Custom note",
                        value=st.session_state.wizard_custom_notes.get(item_key, ""),
                        key=f"custom_{item_key}",
                        height=80,
                        label_visibility="collapsed",
                        placeholder="Add any additional comments specific to this review..."
                    )
                    st.session_state.wizard_custom_notes[item_key] = custom_note
                    
                    st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown("---")
    
    # =========================================================================
    # STEP 3: REVIEW SUMMARY & EXPORT
    # =========================================================================
    st.markdown("---")
    st.subheader("📊 Step 3: Review Summary & Export")
    
    # Summary statistics
    yes_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "Yes")
    no_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "No")
    na_count = sum(1 for v in st.session_state.wizard_checklist_state.values() if v == "N/A")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("✅ Compliant", yes_count)
    with col2:
        st.metric("❌ Issues Found", no_count)
    with col3:
        st.metric("➖ N/A", na_count)
    with col4:
        st.metric("📝 Total Reviewed", yes_count + no_count + na_count)
    
    # Export section
    st.markdown('<div class="export-section">', unsafe_allow_html=True)
    st.markdown("### 📤 Export Review")
    
    if no_count == 0 and (yes_count + na_count) > 0:
        st.success("✅ No issues found! All reviewed items are compliant.")
    elif no_count > 0:
        st.warning(f"⚠️ {no_count} issue(s) found that require comments.")
    
    # -----------------------------------------------------------------
    # Row 1: Word Document + Clear Review
    # -----------------------------------------------------------------
    col1, col2 = st.columns(2)
    
    with col1:
        if DOCX_AVAILABLE:
            if st.button("📄 Generate Word Document", type="primary", use_container_width=True):
                if not st.session_state.wizard_permit_number:
                    st.error("Please enter a permit number before exporting.")
                elif completed_items == 0:
                    st.error("Please review at least one item before exporting.")
                else:
                    doc_buffer = generate_word_document()
                    if doc_buffer:
                        filename = f"Review_{st.session_state.wizard_permit_number}_{datetime.now().strftime('%Y%m%d')}.docx"
                        st.download_button(
                            label="⬇️ Download Word Document",
                            data=doc_buffer,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
        else:
            st.warning("python-docx not available. Install it to enable Word export.")
    
    with col2:
        if st.button("🗑️ Clear Review", use_container_width=True):
            reset_checklist()
            st.session_state.wizard_permit_number = ""
            st.session_state.wizard_address = ""
            st.session_state.wizard_reviewer = None
            st.rerun()

    # -----------------------------------------------------------------
    # Row 2: LAMA CSV + Bluebeam BAX
    # Only shown when there are comments to export (no_count > 0)
    # -----------------------------------------------------------------
    if no_count > 0:
        st.markdown("#### 📊 Extract Comments")

        permit_num = st.session_state.wizard_permit_number or "review"
        datestamp = datetime.now().strftime('%Y%m%d')

        col_e1, col_e2 = st.columns(2)

        with col_e1:
            lama_data = generate_lama_csv()
            if lama_data:
                st.download_button(
                    label="📥 Create CSV File of Comments",
                    data=lama_data,
                    file_name=f"LAMA_Comments_{permit_num}_{datestamp}.csv",
                    mime="text/csv",
                    use_container_width=True,
                    help="Single-column CSV for the LAMA Comment Uploader extension"
                )

        with col_e2:
            bax_data = generate_bluebeam_bax()
            if bax_data:
                st.download_button(
                    label="📐 Create Bluebeam Comments File",
                    data=bax_data,
                    file_name=f"Markups_{permit_num}_{datestamp}.bax",
                    mime="application/octet-stream",
                    use_container_width=True,
                    help="Import into Bluebeam via Markup → Import (.bax format with full styling)"
                )
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Quick copy section for comments
    if no_count > 0:
        st.markdown("---")
        st.subheader("📋 Quick Copy - All Comments")
        st.caption("Copy these comments directly into Bluebeam or your permit system:")
        
        all_comments = []
        for section_id, section_data in checklist.items():
            for item in section_data["items"]:
                item_key = item["id"]
                if st.session_state.wizard_checklist_state.get(item_key) == "No":
                    selected = st.session_state.wizard_selected_comments.get(item_key, [])
                    custom_note = st.session_state.wizard_custom_notes.get(item_key, "")
                    
                    for comment_id in selected:
                        comment_text = COMMENTS.get(comment_id, "")
                        if comment_text:
                            all_comments.append(f"[{comment_id}] {comment_text}")
                    
                    if custom_note.strip():
                        all_comments.append(f"[CUSTOM] {custom_note}")
        
        if all_comments:
            comments_text = "\n\n".join(f"{i+1}. {c}" for i, c in enumerate(all_comments))
            st.text_area(
                "All Comments",
                value=comments_text,
                height=300,
                label_visibility="collapsed"
            )
    
    # Navigation (Admin Panel removed - now 2 columns)
    st.markdown("---")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🏠 Home"):
            st.switch_page("app.py")
    with col2:
        if st.button("💬 Q&A Mode"):
            st.switch_page("pages/1_QA_Mode.py")


if __name__ == "__main__":
    main()
